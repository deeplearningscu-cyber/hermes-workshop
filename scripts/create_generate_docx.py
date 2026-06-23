from docx import Document

doc = Document()
doc.add_heading('docx-generation 安裝與使用說明（繁體中文）', level=1)
doc.add_paragraph('這份說明將示範如何安裝與執行 docx-generation skill，並自動生成相對應的 .docx 文件。')
output_path = 'D:/workspace/Hermes-UAMS/WorkshopV2/workspace/docs/docx-generation_zh_tw.docx'
doc.save(output_path)