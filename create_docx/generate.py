#!/usr/bin/env python3
"""
generate.py - Core logic for creating a .docx file using python-docx.

Usage (as a library):
    from generate import build_docx
    build_docx(title="Meeting Notes", content="Agenda:\n- Item 1", output="notes.docx")

When executed as a script:
    python generate.py --title "My Doc" --content "Hello world" --output out.docx
"""

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _add_heading(doc: Document, title: str, level: int = 1) -> None:
    """Add a heading of the given level with the supplied title."""
    # python-docx uses style names like 'Heading 1', 'Heading 2', etc.
    heading_styles = {
        0: "Heading 1",
        1: "Heading 2",
        2: "Heading 3",
        3: "Heading 4",
        4: "Heading 5",
        5: "Heading 6",
    }
    style_name = heading_styles.get(level, "Heading 1")
    doc.add_paragraph(title, style=style_name)


def _add_content(doc: Document, content: str) -> None:
    """Add the main content as a paragraph."""
    doc.add_paragraph(content)


def build_docx(title: str, content: str, output: str = "output.docx") -> Path:
    """
    Build a .docx file with a title heading and body content.

    Parameters
    ----------
    title: str
        Document title, added as a Heading 1.
    content: str
        Body text, added as a normal paragraph.
    output: str
        Destination filename. Will be created in the current working directory.

    Returns
    -------
    Path
        Path to the generated .docx file.

    Raises
    ------
    ValueError
        If title or content is empty.
    OSError
        If the file cannot be written.
    """
    if not title.strip():
        raise ValueError("Title cannot be empty.")
    if not content.strip():
        raise ValueError("Content cannot be empty.")

    doc = Document()

    # Title as heading
    _add_heading(doc, title, level=0)

    # Body content
    _add_content(doc, content)

    # Optional: set a default font size for the whole document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)

    out_path = Path(output)
    try:
        doc.save(out_path)
    except Exception as exc:
        raise OSError(f"Failed to save {out_path}: {exc}") from exc

    return out_path


def _cli() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a simple .docx file with a title and body content."
    )
    parser.add_argument("--title", required=True, help="Document title")
    parser.add_argument(
        "--content", required=True, help="Body text (can contain newlines)"
    )
    parser.add_argument(
        "--output",
        default="output.docx",
        help="Output .docx filename (default: output.docx)",
    )
    args = parser.parse_args()

    try:
        generated_path = build_docx(
            title=args.title, content=args.content, output=args.output
        )
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        sys.exit(1)

    print(f"✅ Document generated: {generated_path.resolve()}")


if __name__ == "__main__":
    _cli()