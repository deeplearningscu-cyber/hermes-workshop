from docx import Document
import pathlib

# Path to the markdown file we created earlier
md_path = r"D:\workspace\Hermes-UAMS\WorkshopV2\workspace\docs\docx-generation_zh_tw.md"

# Load the markdown content
doc = Document()

with open(md_path, encoding="utf-8") as f:
    for line in f:
        stripped = line.rstrip("\n")
        # Handle headings
        if stripped.startswith("# "):
            text = stripped[2:]
            doc.add_heading(text, level=1)
        elif stripped.startswith("## "):
            text = stripped[3:]
            doc.add_heading(text, level=2)
        elif stripped.startswith("### "):
            text = stripped[4:]
            doc.add_heading(text, level=3)
        # Handle bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            doc.add_paragraph(text, style="List Bullet")
        # Regular paragraph
        else:
            doc.add_paragraph(stripped)

# Save as .docx with same base name
output_path = pathlib.Path(md_path).with_suffix(".docx")
doc.save(str(output_path))

print(f"Generated {output_path}")